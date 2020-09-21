import json
from json import JSONDecodeError
from zipfile import BadZipFile

from django.db import IntegrityError
from django.http import JsonResponse
from django.utils.datastructures import MultiValueDictKeyError
from django.utils.decorators import method_decorator
from django.views import View
from docx import Document

from authorization.decorator import permission_needed
from novel.decorator import get_novel_by_path
from novel.models import Novel, Introduction


class GetAllNovels(View):
    def get(self, request, *args, **kwargs):
        langs = Novel._meta.get_field('lang').choices
        resp = {}
        for lang, _ in langs:
            resp[lang] = []
            novels = Novel.objects.filter(lang=lang, private=False).order_by('-uploadedAt')
            yr = -1
            for n in novels:
                if yr != n.uploadedAt.year:
                    yr = n.uploadedAt.year
                    resp[lang].append({})
                    resp[lang][-1][yr] = []
                resp[lang][-1][yr].append({
                    'title': n.title,
                    'path': n.path,
                    'lore': n.lore,
                    'uploadedAt': n.uploadedAt.isoformat()
                })

        return JsonResponse(resp, charset='utf-8')


class NovelTools(View):
    @method_decorator(get_novel_by_path)
    def get(self, request, *args, **kwargs):
        novel = request.novel
        return JsonResponse({
            'title': novel.title,
            'lore': novel.lore,
            'content': novel.content,
            'uploadedAt': novel.uploadedAt.isoformat(),
            'favorite': request.fb_user.favorites.filter(pk=novel.pk).exists(),
            'lang': novel.lang
        }, charset='utf-8')

    @method_decorator(get_novel_by_path)
    @method_decorator(
        permission_needed(lambda request: not request.fb_user.isAdmin, 'You have to be logged in to edit novels',
                          'You don\'t have permission to edit this novel'))
    def patch(self, request, *args, **kwargs):
        fields = {'title': str, 'lore': str, 'content': str, 'lang': str, 'private': bool}
        novel = request.novel
        try:
            body = json.loads(request.body.decode('utf-8'))
        except JSONDecodeError:
            return JsonResponse({'error': 'Json parse error'}, status=400)
        for p in fields.keys():
            if p in body and type(body[p]) == fields[p] and (p != 'lang' or p in Novel._meta.get_field('lang').choices):
                setattr(novel, p, body[p])
        novel.save()
        return JsonResponse({
            'title': novel.title,
            'path': novel.path,
            'lore': novel.lore,
            'content': novel.content,
            'lang': novel.lang,
            'private': novel.private
        })

    @method_decorator(get_novel_by_path)
    @method_decorator(
        permission_needed(lambda request: not request.fb_user.isAdmin, 'You have to be logged in to edit novels',
                          'You don\'t have permission to edit this novel'))
    def delete(self, request, *args, **kwargs):
        request.novel.delete()
        return JsonResponse({"success": "Deleted successfully"})


class NovelFavoriteToggle(View):
    @method_decorator(get_novel_by_path)
    @method_decorator(
        permission_needed(lambda request: request.fb_user.isAnonymous, 'Log in to mark novels as favorite',
                          "Log in with a non-Anonymous account"))
    def post(self, request, *args, **kwargs):
        novel = request.novel
        favs = request.fb_user.favorites
        if favs.filter(pk=novel.pk).exists():
            favs.remove(novel)
            return JsonResponse({'favorite': False})
        else:
            favs.add(novel)
            return JsonResponse({'favorite': True})

    @method_decorator(get_novel_by_path)
    @method_decorator(permission_needed(lambda request: not request.fb_user.isAdmin, 'Log in to see this value',
                                        'You have to be admin to see this value'))
    def get(self, request, *args, **kwargs):
        return JsonResponse({'favorites': len(request.novel.user_set.all())})


class UserFavorites(View):
    @method_decorator(
        permission_needed(lambda request: request.fb_user.isAnonymous, 'Log in to have novels as favorite',
                          "Log in with a non-Anonymous account"))
    def get(self, request, *args, **kwargs):
        resp = []
        for fav in request.fb_user.favorites.all():
            resp.append({
                'title': fav.title,
                'path': fav.path
            })
        return JsonResponse(resp, safe=False, charset='utf-8')


class NewUpload(View):
    @method_decorator(
        permission_needed(lambda request: not request.fb_user.isAdmin, 'You have to be logged in to upload novels',
                          'You don\'t have permission to upload novels'))
    def post(self, request, *args, **kwargs):
        fn = 'noveldoc'
        try:
            d = Document(request.FILES[fn])
        except MultiValueDictKeyError:
            return JsonResponse({"error": f"No file with name: {fn}"}, status=400)
        except BadZipFile:
            return JsonResponse({"error": "Wrong file"}, status=400)
        if len(d.paragraphs) < 2:
            return JsonResponse({"error": "Wrong file"}, status=400)
        novel = None
        try:
            novel = Novel.objects.create(title=d.paragraphs[0].text, private=True)
        except IntegrityError:
            n = Novel.objects.get(title=d.paragraphs[0].text)
            if n.private:
                n.delete()
                novel = Novel.objects.create(title=d.paragraphs[0].text, private=True)
            else:
                return JsonResponse({"error": "Not unique title"}, status=400)
        # string.join won't work
        content = ''
        padding = '\n'
        for l in d.paragraphs[1:]:
            content += l.text + padding
        content = content[:-len(padding)]
        novel.content = content
        novel.save()
        return JsonResponse({
            'title': novel.title,
            'path': novel.path,
            'filename': request.FILES[fn].name
        })


class IntroductionView(View):
    @method_decorator(
        permission_needed(lambda request: not request.fb_user.isAdmin, 'You have to be logged in to edit introduction',
                          'You don\'t have permission to edit introduction'))
    def post(self, request, *args, **kwargs):
        try:
            body = json.loads(request.body.decode('utf-8'))
        except JSONDecodeError:
            return JsonResponse({'error': 'Json parse error'}, status=400)
        if not ('introduction' in body and isinstance(body['introduction'], str)):
            return JsonResponse({'error': 'introduction property missing'}, status=400)
        if not ('lang' in body and isinstance(body['introduction'], str)):
            return JsonResponse({'error': 'lang property missing'}, status=400)
        intro, _ = Introduction.objects.update_or_create(key=body['lang'], value=body['introduction'])
        intro.save()
        return JsonResponse({'introduction': intro.value})

    def get(self, request, *args, **kwargs):
        lang: str = request.GET.get('lang', 'EN')
        try:
            intro = Introduction.objects.get(key=lang).value
        except:
            intro = 'Sorry, I don\'t use this language ¯\\_( ͡° ͜ʖ ͡°)_/¯'
        return JsonResponse({'introduction': intro})
